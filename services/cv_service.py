from __future__ import annotations

import pprint
import tempfile
from pathlib import Path

ROOT = Path(__file__).parent.parent
CV_PATH = ROOT / "data" / "master_cv.md"
TEMPLATE_PATH = ROOT / "templates" / "cv_template_uploaded.py"


# ---------------------------------------------------------------------------
# Defaults
# ---------------------------------------------------------------------------

def _default_template() -> dict:
    return {
        "name": "Uploaded",
        "page_margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        "fonts": {
            "name_header": {"name": "Calibri", "size": 22, "bold": True,  "color": "000000"},
            "contact":     {"name": "Calibri", "size": 10, "bold": False, "color": "444444"},
            "heading1":    {"name": "Calibri", "size": 13, "bold": True,  "color": "000000"},
            "heading2":    {"name": "Calibri", "size": 11, "bold": True,  "color": "2E74B5"},
            "body":        {"name": "Calibri", "size": 10, "bold": False, "color": "000000"},
            "bullet":      {"name": "Calibri", "size": 10, "bold": False, "color": "000000"},
            "dates":       {"name": "Calibri", "size": 10, "bold": False, "color": "666666"},
        },
        "spacing": {
            "before_section_pt": 10,
            "after_section_pt":  4,
            "line_spacing":      1.15,
            "bullet_indent_cm":  0.5,
        },
        "section_divider": True,
        "bullet_char": "•",
    }


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _docx_to_markdown(doc) -> str:
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style_name = (para.style.name or "").lower()

        if "heading 1" in style_name:
            lines.append(f"## {text}")
        elif "heading 2" in style_name:
            lines.append(f"### {text}")
        elif "list bullet" in style_name or "list number" in style_name:
            lines.append(f"- {text}")
        else:
            parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    parts.append(f"***{t}***")
                elif run.bold:
                    parts.append(f"**{t}**")
                elif run.italic:
                    parts.append(f"*{t}*")
                else:
                    parts.append(t)
            lines.append("".join(parts) if parts else text)

    return "\n".join(lines)


def _detect_template_from_docx(doc) -> dict:
    template = _default_template()

    # Margins
    try:
        section = doc.sections[0]
        if section.top_margin:
            template["page_margins"]["top"] = round(section.top_margin.inches, 2)
        if section.bottom_margin:
            template["page_margins"]["bottom"] = round(section.bottom_margin.inches, 2)
        if section.left_margin:
            template["page_margins"]["left"] = round(section.left_margin.inches, 2)
        if section.right_margin:
            template["page_margins"]["right"] = round(section.right_margin.inches, 2)
    except Exception:
        pass

    name_done = heading1_done = heading2_done = body_done = False

    for para in doc.paragraphs:
        if not para.runs or not para.text.strip():
            continue

        run = para.runs[0]
        font_name = run.font.name or "Calibri"

        try:
            font_size = round(run.font.size.pt) if run.font.size else None
        except Exception:
            font_size = None

        try:
            rgb = run.font.color.rgb if run.font.color and run.font.color.type else None
            color_hex = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}" if rgb else None
        except Exception:
            color_hex = None

        is_bold = bool(run.font.bold)
        style_name = (para.style.name or "").lower()

        def make(size_default: int, bold_default: bool, color_default: str) -> dict:
            return {
                "name": font_name,
                "size": font_size or size_default,
                "bold": is_bold if run.font.bold is not None else bold_default,
                "color": color_hex or color_default,
            }

        if not name_done and font_size and font_size >= 16:
            template["fonts"]["name_header"] = make(22, True, "000000")
            name_done = True
        elif not heading1_done and (
            "heading 1" in style_name or (font_size and font_size >= 12 and is_bold)
        ):
            template["fonts"]["heading1"] = make(13, True, "000000")
            heading1_done = True
        elif not heading2_done and (
            "heading 2" in style_name or (font_size and font_size >= 11 and is_bold)
        ):
            template["fonts"]["heading2"] = make(11, True, "2E74B5")
            heading2_done = True
        elif not body_done and not is_bold:
            template["fonts"]["body"] = make(10, False, "000000")
            template["fonts"]["bullet"] = make(10, False, "000000")
            body_done = True

        if name_done and heading1_done and heading2_done and body_done:
            break

    return template


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _pdf_to_markdown_and_template(pdf_path: Path) -> tuple[str, dict]:
    import fitz  # pymupdf

    doc = fitz.open(str(pdf_path))
    all_spans: list[dict] = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if text:
                        all_spans.append({
                            "text": text,
                            "size": round(span["size"]),
                            "font": span["font"],
                            "color": span["color"],
                            "flags": span["flags"],
                            "bbox": span["bbox"],
                        })
    doc.close()

    if not all_spans:
        return "", _default_template()

    sizes = sorted(set(s["size"] for s in all_spans), reverse=True)
    name_size = sizes[0] if sizes else 18
    h1_size = sizes[1] if len(sizes) > 1 else 13
    h2_size = sizes[2] if len(sizes) > 2 else 11

    lines: list[str] = []
    prev_bottom: float | None = None

    for span in all_spans:
        text = span["text"]
        size = span["size"]
        is_bold = bool(span["flags"] & 16)
        top = span["bbox"][1]

        if prev_bottom is not None and top > prev_bottom + 4:
            lines.append("")

        if size >= name_size:
            lines.append(f"# {text}")
        elif size >= h1_size and is_bold:
            lines.append(f"## {text}")
        elif size >= h2_size and is_bold:
            lines.append(f"### {text}")
        elif text[:1] in ("•", "·", "▪", "◦"):
            lines.append(f"- {text[1:].strip()}")
        else:
            lines.append(text)

        prev_bottom = span["bbox"][3]

    md = "\n".join(lines)

    # Build template from detected fonts
    def _color_int_to_hex(c: int) -> str:
        return f"{(c >> 16) & 0xFF:02X}{(c >> 8) & 0xFF:02X}{c & 0xFF:02X}"

    def _span_font(span: dict | None, size_default: int, bold: bool, color_default: str) -> dict:
        if not span:
            return {"name": "Helvetica", "size": size_default, "bold": bold, "color": color_default}
        raw_font = span["font"].split(",")[0].split("-")[0].strip() or "Helvetica"
        return {
            "name": raw_font,
            "size": span["size"],
            "bold": bold,
            "color": _color_int_to_hex(span["color"]) if span["color"] else color_default,
        }

    name_span = next((s for s in all_spans if s["size"] >= name_size), None)
    h1_span = next((s for s in all_spans if s["size"] >= h1_size and bool(s["flags"] & 16)), None)
    body_span = next(
        (s for s in reversed(all_spans) if not bool(s["flags"] & 16)),
        None,
    )

    template = _default_template()
    template["fonts"]["name_header"] = _span_font(name_span, 22, True, "000000")
    template["fonts"]["heading1"]    = _span_font(h1_span,   13, True, "000000")
    template["fonts"]["body"]        = _span_font(body_span, 10, False, "000000")
    template["fonts"]["bullet"]      = _span_font(body_span, 10, False, "000000")

    return md, template


# ---------------------------------------------------------------------------
# Persistence
# ---------------------------------------------------------------------------

def _save_template(template: dict) -> None:
    content = f"TEMPLATE = {pprint.pformat(template, indent=4, width=80)}\n"
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    TEMPLATE_PATH.write_text(content, encoding="utf-8")


def _update_env_cv_template(slug: str) -> None:
    env_path = ROOT / ".env"
    existing = env_path.read_text(encoding="utf-8").splitlines() if env_path.exists() else []
    filtered = [l for l in existing if not l.startswith("CV_TEMPLATE=")]
    filtered.append(f"CV_TEMPLATE={slug}")
    env_path.write_text("\n".join(filtered) + "\n", encoding="utf-8")
    from config import reset_settings
    reset_settings()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def process_upload(file_bytes: bytes, filename: str) -> dict:
    suffix = Path(filename).suffix.lower()

    if suffix == ".docx":
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        md_content = _docx_to_markdown(doc)
        template = _detect_template_from_docx(doc)

    elif suffix == ".pdf":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = Path(tmp.name)
        try:
            md_content, template = _pdf_to_markdown_and_template(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)

    else:
        raise ValueError(f"Unsupported file type '{suffix}'. Upload a .pdf or .docx file.")

    if not md_content.strip():
        raise ValueError("Could not extract any text from the uploaded file.")

    CV_PATH.parent.mkdir(parents=True, exist_ok=True)
    CV_PATH.write_text(md_content, encoding="utf-8")

    _save_template(template)
    _update_env_cv_template("uploaded")

    return {
        "chars": len(md_content),
        "template_name": template["name"],
        "margins": template["page_margins"],
        "fonts_detected": {k: v["name"] for k, v in template["fonts"].items()},
    }


def get_cv_content() -> str:
    if not CV_PATH.exists():
        return ""
    return CV_PATH.read_text(encoding="utf-8")
