from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from generators.template_engine import TemplateEngine

ROOT = Path(__file__).parent.parent
RESUMES_DIR = ROOT / "resumes" / "tailored"
CL_DIR = ROOT / "outputs" / "cover_letters"

_engine = TemplateEngine()


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^\w\s-]", "", text).strip()
    return re.sub(r"[\s]+", "_", text)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_font(run, font_def: dict) -> None:
    run.font.name = font_def.get("name", "Calibri")
    run.font.size = Pt(font_def.get("size", 10))
    run.font.bold = font_def.get("bold", False)
    run.font.color.rgb = _hex_to_rgb(font_def.get("color", "000000"))


def _add_hyperlink_run(paragraph, text: str, url: str, font_def: dict) -> None:
    """Append a clickable hyperlink run to an existing paragraph via XML."""
    import docx.opc.constants as _opc

    part = paragraph.part
    r_id = part.relate_to(url, _opc.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_def.get("name", "Calibri"))
    rFonts.set(qn("w:hAnsi"), font_def.get("name", "Calibri"))
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_def.get("size", 10) * 2)))
    rPr.append(sz)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), font_def.get("color", "333333"))
    rPr.append(color_el)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def _add_bottom_border(p, color: str = "CCCCCC") -> None:
    """Add a bottom border to a paragraph (used as section divider)."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, text: str, fonts: dict, spacing: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing.get("before_section_pt", 10))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _apply_font(run, fonts["heading1"])
    _add_bottom_border(p)


class DocxBuilder:
    def build_cv(
        self,
        tailored_data: dict,
        person_name: str,
        company: str,
        template_name: str = "professional",
        person_email: str = "",
        person_phone: str = "",
        person_address: str = "",
        person_linkedin: str = "",
        outputs_root: Path | None = None,
    ) -> Path:
        template = _engine.load(template_name)
        fonts   = template["fonts"]
        spacing = template["spacing"]
        doc = Document()
        _engine.apply_margins(doc, template)

        # ── Name ──────────────────────────────────────────────────────────────
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(person_name)
        _apply_font(name_run, fonts["name_header"])

        # ── Contact row: email | phone | address ──────────────────────────────
        contact_parts = [p for p in [person_email, person_phone, person_address] if p.strip()]
        if contact_parts:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(1)
            _apply_font(cp.add_run("   |   ".join(contact_parts)), fonts["contact"])

        # ── LinkedIn row: [in] badge + clickable URL ──────────────────────────
        if person_linkedin.strip():
            li_raw = person_linkedin.strip()
            li_url = li_raw if li_raw.startswith("http") else "https://" + li_raw
            li_display = re.sub(r"^https?://", "", li_raw).rstrip("/")

            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.paragraph_format.space_after = Pt(3)

            # Badge: white "in" on LinkedIn-blue background (brand icon — stays blue)
            badge_run = lp.add_run(" in ")
            _apply_font(badge_run, {"name": fonts["contact"]["name"], "size": 8, "bold": True, "color": "FFFFFF"})
            rPr = badge_run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "0077B5")
            rPr.append(shd)

            lp.add_run("  ")

            # Clickable URL in dark grey (no blue in the CV text)
            _add_hyperlink_run(lp, li_display, li_url, {**fonts["contact"], "color": "333333"})

        # Separator line under header
        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(2)
        sep_p.paragraph_format.space_after = Pt(4)
        _add_bottom_border(sep_p, color="CCCCCC")

        # ── Summary ───────────────────────────────────────────────────────────
        summary: str = tailored_data.get("summary", "")
        if summary:
            _add_section_heading(doc, "Professional Summary", fonts, spacing)
            for line in summary.split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    _apply_font(p.add_run(line), fonts["body"])

        # ── Skills ────────────────────────────────────────────────────────────
        skills: dict = tailored_data.get("skills", {})
        if any(v for v in skills.values()):
            _add_section_heading(doc, "Skills", fonts, spacing)
            label_map = {
                "languages":  "Languages",
                "frameworks": "Frameworks & Libraries",
                "databases":  "Databases",
                "cloud_infra": "Cloud & Infrastructure",
                "tools":      "Tools & Practices",
                "other":      "Other",
            }
            for key, items in skills.items():
                if not items:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                label = label_map.get(key, key.replace("_", " ").capitalize())
                # Label: slightly darker than body, no bold, so values don't inherit highlight
                _apply_font(p.add_run(f"{label}: "), {**fonts["body"], "color": "555555"})
                _apply_font(p.add_run(", ".join(items)), fonts["body"])

        # ── Experience ────────────────────────────────────────────────────────
        experience = tailored_data.get("experience", [])
        if experience:
            _add_section_heading(doc, "Experience", fonts, spacing)
            for entry in experience:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                _apply_font(
                    p.add_run(f"{entry.get('title', '')}  —  {entry.get('company', '')}"),
                    fonts["heading2"],
                )
                meta = f"{entry.get('dates', '')}  |  {entry.get('location', '')}".strip(" |")
                if meta:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_after = Pt(2)
                    _apply_font(p2.add_run(meta), fonts["dates"])
                for bullet in entry.get("bullets", []):
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.left_indent = Inches(spacing.get("bullet_indent_cm", 0.5) / 2.54)
                    bp.paragraph_format.space_after = Pt(1)
                    _apply_font(bp.add_run(bullet.lstrip("•- ").strip()), fonts["bullet"])

        # ── Certifications ────────────────────────────────────────────────────
        certifications = tailored_data.get("certifications", [])
        if certifications:
            _add_section_heading(doc, "Certifications", fonts, spacing)
            for cert in certifications:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                name_str = cert.get("name", "")
                issuer   = cert.get("issuer", "")
                year     = cert.get("year", "")
                _apply_font(p.add_run(name_str), {**fonts["body"], "bold": True})
                suffix_parts = [s for s in [issuer, year] if s]
                if suffix_parts:
                    _apply_font(p.add_run(f"  —  {',  '.join(suffix_parts)}"), fonts["dates"])

        # ── Education ─────────────────────────────────────────────────────────
        education = tailored_data.get("education", [])
        if education:
            _add_section_heading(doc, "Education", fonts, spacing)
            for edu in education:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                _apply_font(
                    p.add_run(f"{edu.get('degree', '')}  —  {edu.get('institution', '')}"),
                    fonts["heading2"],
                )
                year_note = f"{edu.get('year', '')}  {edu.get('notes', '')}".strip()
                if year_note:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_after = Pt(2)
                    _apply_font(p2.add_run(year_note), fonts["dates"])

        if outputs_root is not None:
            out_dir = outputs_root / _safe_filename(company)
        else:
            out_dir = RESUMES_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{_safe_filename(person_name)}_cv.docx"
        doc.save(out_path)
        return out_path

    def build_cover_letter(
        self,
        content: str,
        person_name: str,
        company: str,
        template_name: str = "professional",
    ) -> Path:
        CL_DIR.mkdir(parents=True, exist_ok=True)
        template = _engine.load(template_name)
        fonts = template["fonts"]
        doc = Document()
        _engine.apply_margins(doc, template)

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _apply_font(p.add_run(line), fonts["body"])

        out_path = CL_DIR / f"Cover_Letter_{_safe_filename(person_name)}_{_safe_filename(company)}.docx"
        doc.save(out_path)
        return out_path
